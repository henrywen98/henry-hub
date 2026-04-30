"""Pytest configuration for doc-feature-extractor.

把 scripts/ 加进 sys.path,让 test 文件可以 `from extract_solution import ...`。
所有 fixtures 都构造合成数据 (无客户真实文档),便于在 CI / 沙箱环境跑。
"""
from __future__ import annotations

import sys
from pathlib import Path

# 让 tests/ 文件能 import scripts/ 下的模块
SCRIPTS_DIR = Path(__file__).resolve().parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook


# ---------------------------------------------------------------------------
# 合成 docx (Mode B 用)
# ---------------------------------------------------------------------------


@pytest.fixture
def synthetic_solution_docx(tmp_path: Path) -> Path:
    """合成解决方案 docx: H1 大标题 + H3 模块 (5 个) + 内嵌"功能模块/费用/工作量"表。

    模块层级是 H3 (count=5,触发 detect_heading_level 主路径)。
    """
    doc = DocxDocument()
    doc.add_heading("合成方案 V1", level=1)
    doc.add_paragraph("这是合成的测试方案,用于验证 extract_solution 行为。")

    doc.add_heading("第一章 业务背景", level=2)
    doc.add_paragraph("业务背景描述。")

    doc.add_heading("第二章 功能模块", level=2)

    modules = [
        ("项目智能分析", "通过 AI 抽取项目核心信息并生成画像。"),
        ("合同智能审查", "对合同条款做合规检查与风险提示。"),
        ("财报智能解析", "从 PDF 财报抽取结构化字段。"),
        ("访谈纪要解析", "从 Word/音频访谈记录生成结构化纪要。"),
        ("BP 文件解析", "从 PPT/PDF/Word 格式 BP 抽取字段。"),
    ]
    for title, desc in modules:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    # 内嵌"功能模块 + 费用 + 工作量"表 (5 行,对应 5 个模块)
    doc.add_heading("第三章 费用与工作量", level=2)
    table = doc.add_table(rows=6, cols=3)
    table.rows[0].cells[0].text = "功能模块"
    table.rows[0].cells[1].text = "费用"
    table.rows[0].cells[2].text = "工作量"
    for i, (title, _) in enumerate(modules, 1):
        table.rows[i].cells[0].text = title
        table.rows[i].cells[1].text = f"{i}.5 万"
        table.rows[i].cells[2].text = f"{i*5} 人天"

    out = tmp_path / "synthetic_solution.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def synthetic_solution_no_modules_docx(tmp_path: Path) -> Path:
    """合成 docx 只有正文段落,没有任何 heading → extract 应该 0 模块。"""
    doc = DocxDocument()
    doc.add_paragraph("纯段落,没有标题。")
    doc.add_paragraph("再来一段。")
    out = tmp_path / "no_modules.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def corrupt_docx(tmp_path: Path) -> Path:
    """0 字节伪 docx。"""
    out = tmp_path / "corrupt.docx"
    out.write_bytes(b"")
    return out


# ---------------------------------------------------------------------------
# 合成 xlsx (Mode C 用)
# ---------------------------------------------------------------------------


@pytest.fixture
def synthetic_quotation_xlsx(tmp_path: Path) -> Path:
    """合成报价 xlsx: 2 sheet,典型表头 + 父值继承 + 合计行。"""
    wb = Workbook()
    default = wb.active
    if default is not None:
        wb.remove(default)

    # Sheet 1: 标准报价表 — 业务范围 用 merged cell 模拟 (实际写法是只在第一行写)
    ws1 = wb.create_sheet("客户A-报价清单")
    ws1.append(["业务范围", "模块", "子模块", "功能描述", "分摊报价", "优先级", "备注"])
    ws1.append(["投资管理", "BP 解析", "PDF 解析", "支持 PDF 格式 BP", "5 万", "P0", ""])
    ws1.append(["", "BP 解析", "PPT 解析", "支持 PPT 格式 BP", "3 万", "P1", ""])  # 父值继承
    ws1.append(["", "合同审查", "条款合规", "条款合规检查", "8 万", "P0", "源备注 A"])
    ws1.append(["", "合计", "", "", "16 万", "", ""])  # 合计行 → 必须跳过

    # Sheet 2: 招标功能清单 (无报价列)
    ws2 = wb.create_sheet("客户B-招标功能")
    ws2.append(["业务范围", "功能模块", "二级功能", "功能描述", "需求来源"])
    ws2.append(["风控", "尽调辅助", "风险报告生成", "AI 自动生成尽调报告", "招标"])
    ws2.append(["", "财报识别", "OCR 识别", "财报 OCR 识别", "招标"])

    out = tmp_path / "synthetic_quotation.xlsx"
    wb.save(out)
    return out


@pytest.fixture
def synthetic_quotation_ambiguous_headers_xlsx(tmp_path: Path) -> Path:
    """专门测 T2.1 修复: header [功能模块, 功能说明, 备注说明] → desc 跟 note 不能错位。"""
    wb = Workbook()
    default = wb.active
    if default is not None:
        wb.remove(default)
    ws = wb.create_sheet("test")
    ws.append(["业务范围", "功能模块", "功能说明", "备注说明"])
    ws.append(["A", "模块1", "这是功能描述", "这是备注内容"])
    out = tmp_path / "ambiguous.xlsx"
    wb.save(out)
    return out


@pytest.fixture
def synthetic_no_header_xlsx(tmp_path: Path) -> Path:
    """无可识别表头的 xlsx → extract 应返回 0 行,exit 1。"""
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Sheet1"
    ws["A1"] = "random"
    ws["B1"] = "no headers here"
    out = tmp_path / "no_header.xlsx"
    wb.save(str(out))
    return out
