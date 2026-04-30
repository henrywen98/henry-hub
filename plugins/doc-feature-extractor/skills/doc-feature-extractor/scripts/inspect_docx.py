"""单文件调试: 列出指定 .docx 的所有表格表头与章节标题。"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def inspect(path: Path) -> None:
    doc = Document(str(path))

    print(f"\n=== {path.name} ===\n")

    print("# 章节标题 (Heading 1/2/3)")
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading") and p.text.strip():
            print(f"  [{style}] {p.text.strip()}")

    print(f"\n# 表格 ({len(doc.tables)} 个)")
    for i, table in enumerate(doc.tables, 1):
        if not table.rows:
            continue
        header = [c.text.strip().replace("\n", " ") for c in table.rows[0].cells]
        body_rows = len(table.rows) - 1
        print(f"  [表 {i}] 表头: {header}  (数据行={body_rows})")
        for r in table.rows[1 : 1 + min(2, body_rows)]:
            sample = [c.text.strip().replace("\n", " ")[:40] for c in r.cells]
            print(f"      样例: {sample}")


def main() -> int:
    if len(sys.argv) < 2:
        print("用法: inspect_docx.py <docx 路径> [<docx 路径> ...]", file=sys.stderr)
        return 1
    for arg in sys.argv[1:]:
        p = Path(arg)
        if not p.exists():
            print(f"找不到: {p}", file=sys.stderr)
            continue
        inspect(p)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
