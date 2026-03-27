#!/usr/bin/env python3
"""
从参考文档抽取基线格式，系统性重建 reference.docx 模板。
所有样式定义集中在此，不再零散修补。

基线来源：2-12金财投资企业信息门户升级项目概要设计.docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ══════════════════════════════════════════════
# 基线配置（从参考文档提取 + 段落级实际渲染值）
# ══════════════════════════════════════════════

PAGE_LAYOUT = {
    'cover': {
        'top': Cm(2.54), 'bottom': Cm(2.54),
        'left': Cm(3.17), 'right': Cm(3.17),
        'header_distance': Cm(1.50), 'footer_distance': Cm(1.75),
    },
    'body': {
        'top': Cm(3.50), 'bottom': Cm(2.80),
        'left': Cm(2.80), 'right': Cm(2.60),
        'header_distance': Cm(1.50), 'footer_distance': Cm(1.75),
    },
}

# 样式定义：用段落级实际渲染值（而非原始样式级值）
# Heading 1/2 参考文档样式级是 2.4/1.73，但段落级全部覆盖为 1.5
# Normal 参考文档无行距设置，段落级覆盖为 1.5
STYLE_DEFS = {
    'Normal': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14, 'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'line_spacing': 1.5,
        'color': '000000',
    },
    'Heading 1': {
        'eastAsia': '黑体', 'ascii': '黑体',
        'size_pt': 16, 'bold': None,  # 黑体本身已粗
        'line_spacing': 1.5,  # 段落级实际值
        'space_before_pt': 13, 'space_after_pt': 13,
        'color': '000000',
    },
    'Heading 2': {
        'eastAsia': '楷体', 'ascii': '楷体',
        'size_pt': 16, 'bold': None,
        'line_spacing': 1.5,  # 段落级实际值
        'space_before_pt': 13, 'space_after_pt': 13,
        'color': '000000',
    },
    'Heading 3': {
        'eastAsia': '仿宋', 'ascii': '仿宋',
        'size_pt': 16, 'bold': True,
        'line_spacing': 1.7333,  # 参考文档无段落级覆盖，保持原样式值
        'space_before_pt': 13, 'space_after_pt': 13,
        'color': '000000',
    },
    'Heading 4': {
        'eastAsia': '黑体', 'ascii': 'Times New Roman',
        'size_pt': 14, 'bold': True,
        'line_spacing': 1.5,
        'space_before_pt': 13, 'space_after_pt': 13,
        'color': '000000',
    },
    # pandoc 用 Body Text 渲染正文后续段落
    'Body Text': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14, 'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'line_spacing': 1.5,
        'first_line_indent_cm': 0.74,
        'space_before_pt': 0, 'space_after_pt': 0,
        'color': '000000',
    },
    # pandoc 用 First Paragraph 渲染标题后的首段
    'First Paragraph': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14, 'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'line_spacing': 1.5,
        'first_line_indent_cm': 0.74,
        'color': '000000',
    },
    'Body Text 2': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'line_spacing': 1.5,
        'color': '000000',
    },
    'Body Text 3': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'color': '000000',
    },
    # pandoc 紧凑列表
    'Compact': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'line_spacing': 1.5,
        'space_before_pt': 0, 'space_after_pt': 0,
        'color': '000000',
    },
    'List Paragraph': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 14,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'line_spacing': 1.5,
        'first_line_indent_cm': 0.74,
        'color': '000000',
    },
    # 代码块
    'Source Code': {
        'eastAsia': '仿宋', 'ascii': 'Courier New',
        'size_pt': 10,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'line_spacing': 1.0,
        'space_before_pt': 6, 'space_after_pt': 6,
        'color': '000000',
    },
    # 参考文档原始样式
    'Normal (Web)': {
        'eastAsia': '宋体', 'ascii': '宋体',
        'size_pt': 12,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'space_before_pt': 5, 'space_after_pt': 5,
        'color': '000000',
    },
    # TOC 样式
    'TOC Heading': {
        'eastAsia': '黑体', 'ascii': 'Times New Roman',
        'size_pt': 14, 'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'line_spacing': 1.5,
        'space_before_pt': 12, 'space_after_pt': 6,
        'color': '000000',
    },
    'toc 1': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 12, 'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'space_before_pt': 6, 'space_after_pt': 6,
        'color': '000000',
    },
    'toc 2': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 12,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'color': '000000',
    },
    'toc 3': {
        'eastAsia': '仿宋', 'ascii': 'Times New Roman',
        'size_pt': 12,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'color': '000000',
    },
    # 图片居中
    '图片样式': {
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'color': '000000',
    },
}


def apply_style(style, defs):
    """将样式定义应用到 style 对象"""
    # 字体
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")} />')
        style.element.insert(0, rpr)

    # eastAsia + ascii 字体
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)

    if 'eastAsia' in defs:
        rfonts.set(qn('w:eastAsia'), defs['eastAsia'])
    if 'ascii' in defs:
        rfonts.set(qn('w:ascii'), defs['ascii'])
        rfonts.set(qn('w:hAnsi'), defs['ascii'])

    # 字号
    if 'size_pt' in defs:
        style.font.size = Pt(defs['size_pt'])

    # 加粗
    if 'bold' in defs and defs['bold'] is not None:
        style.font.bold = defs['bold']

    # 颜色
    if 'color' in defs:
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = parse_xml(f'<w:color {nsdecls("w")} />')
            rpr.append(color_el)
        color_el.set(qn('w:val'), defs['color'])
        # 清除 themeColor
        for attr in [qn('w:themeColor'), qn('w:themeTint'), qn('w:themeShade')]:
            if attr in color_el.attrib:
                del color_el.attrib[attr]

    # 段落格式
    pf = style.paragraph_format
    if 'alignment' in defs:
        pf.alignment = defs['alignment']
    if 'line_spacing' in defs:
        pf.line_spacing = defs['line_spacing']
    if 'space_before_pt' in defs:
        pf.space_before = Pt(defs['space_before_pt'])
    if 'space_after_pt' in defs:
        pf.space_after = Pt(defs['space_after_pt'])
    if 'first_line_indent_cm' in defs:
        pf.first_line_indent = Cm(defs['first_line_indent_cm'])


def clean_all_style_colors(doc):
    """清理所有样式中的非黑色字体颜色"""
    for style in doc.styles:
        for color in style.element.findall('.//' + qn('w:color')):
            val = color.get(qn('w:val'))
            if val and val not in ('000000', 'auto'):
                color.set(qn('w:val'), '000000')
                for attr in [qn('w:themeColor'), qn('w:themeTint'), qn('w:themeShade')]:
                    if attr in color.attrib:
                        del color.attrib[attr]


def build_template(src_path, dst_path):
    """从源文档构建模板"""
    # 复制源文档（保留完整的 styles.xml 结构）
    doc = Document(src_path)

    # 清空内容
    body = doc.element.body
    while len(doc.paragraphs) > 1:
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
    if doc.paragraphs:
        doc.paragraphs[0].clear()
    while doc.tables:
        t = doc.tables[-1]
        t._element.getparent().remove(t._element)

    # 1. 清理所有样式颜色 + 移除 Heading 自动编号
    clean_all_style_colors(doc)
    for hname in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
        try:
            hs = doc.styles[hname]
            ppr = hs.element.find(qn('w:pPr'))
            if ppr is not None:
                numpr = ppr.find(qn('w:numPr'))
                if numpr is not None:
                    ppr.remove(numpr)
        except KeyError:
            pass

    # 2. 应用基线样式定义
    for style_name, defs in STYLE_DEFS.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            # 创建缺失的样式
            style = doc.styles.add_style(style_name, 1)  # paragraph
            style.base_style = doc.styles['Normal']
        apply_style(style, defs)

    # 3. 页面设置（使用正文节的边距）
    for section in doc.sections:
        layout = PAGE_LAYOUT['body']
        section.top_margin = layout['top']
        section.bottom_margin = layout['bottom']
        section.left_margin = layout['left']
        section.right_margin = layout['right']
        section.header_distance = layout['header_distance']
        section.footer_distance = layout['footer_distance']

    # 4. 清空页眉页脚
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            p.clear()
        for p in sec.footer.paragraphs:
            p.clear()

    doc.save(dst_path)


def verify_template(path):
    """验证模板"""
    doc = Document(path)
    print(f"\n═══ 模板验证: {path} ═══\n")

    issues = []
    for style_name, defs in STYLE_DEFS.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            issues.append(f"样式 {style_name} 不存在")
            continue

        # 检查 eastAsia
        rpr = style.element.find(qn('w:rPr'))
        if rpr is not None and 'eastAsia' in defs:
            rf = rpr.find(qn('w:rFonts'))
            ea = rf.get(qn('w:eastAsia'), '') if rf is not None else ''
            if ea != defs['eastAsia']:
                issues.append(f"{style_name} eastAsia: {ea} → {defs['eastAsia']}")

        # 检查颜色
        if rpr is not None:
            c = rpr.find(qn('w:color'))
            if c is not None:
                val = c.get(qn('w:val'))
                if val and val not in ('000000', 'auto'):
                    issues.append(f"{style_name} color: #{val}")

        # 检查行距
        if 'line_spacing' in defs:
            ls = style.paragraph_format.line_spacing
            if ls is None or abs(ls - defs['line_spacing']) > 0.05:
                issues.append(f"{style_name} line_spacing: {ls} → {defs['line_spacing']}")

        print(f"  ✓ {style_name}")

    if issues:
        print(f"\n  问题:")
        for iss in issues:
            print(f"    ✗ {iss}")
    else:
        print(f"\n  全部通过 ({len(STYLE_DEFS)} 个样式)")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: build_template.py <源docx文件> [输出路径]")
        print("示例: build_template.py 参考概设.docx")
        sys.exit(1)
    SRC = sys.argv[1]
    SCRIPT_DIR = Path(__file__).resolve().parent
    DST = sys.argv[2] if len(sys.argv) > 2 else str(SCRIPT_DIR.parent / "templates" / "reference.docx")

    print(f"构建模板: {DST}")
    build_template(SRC, DST)
    verify_template(DST)
    print(f"\n模板已保存到: {DST}")
