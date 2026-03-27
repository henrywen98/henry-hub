#!/usr/bin/env python3
"""
pandoc 生成的 docx 后处理脚本：
- 插入封面页（标题、副标题、公司、日期）
- 插入变更记录空表
- 修正中文字体（run 级别确保字体生效）
- 设置页眉
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# 样式 → (中文字体, 西文字体)
FONT_MAP = {
    'Heading 1': ('黑体', 'Times New Roman'),
    'Heading 2': ('楷体', 'Times New Roman'),
    'Heading 3': ('仿宋', 'Times New Roman'),
    'Heading 4': ('黑体', 'Times New Roman'),
    'Normal': ('仿宋', 'Times New Roman'),
    'Body Text': ('仿宋', 'Times New Roman'),
    'Body Text 2': ('仿宋', 'Times New Roman'),
    'Body Text 3': ('仿宋', 'Times New Roman'),
    'First Paragraph': ('仿宋', 'Times New Roman'),
    'Compact': ('仿宋', 'Times New Roman'),
    'List Paragraph': ('仿宋', 'Times New Roman'),
    'Source Code': ('仿宋', 'Courier New'),
}


def make_para(text='', cn_font='仿宋', en_font='Times New Roman',
              size=28, bold=False, align='center'):
    """生成一个带完整字体设置的段落 XML"""
    bold_xml = '<w:b/>' if bold else ''
    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:jc w:val="{align}"/></w:pPr>'
    )
    if text:
        p_xml += (
            f'  <w:r>'
            f'    <w:rPr>'
            f'      <w:rFonts w:eastAsia="{cn_font}" w:ascii="{en_font}" w:hAnsi="{en_font}"/>'
            f'      {bold_xml}'
            f'      <w:sz w:val="{size}"/><w:szCs w:val="{size}"/>'
            f'    </w:rPr>'
            f'    <w:t xml:space="preserve">{text}</w:t>'
            f'  </w:r>'
        )
    p_xml += '</w:p>'
    return parse_xml(p_xml)


def make_section_break(top=1440, bottom=1440, left=1800, right=1800):
    """生成分节符段落"""
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:sectPr>'
        f'      <w:pgSz w:w="11906" w:h="16838"/>'
        f'      <w:pgMar w:top="{top}" w:right="{right}" w:bottom="{bottom}" w:left="{left}" '
        f'              w:header="851" w:footer="992"/>'
        f'    </w:sectPr>'
        f'  </w:pPr>'
        f'</w:p>'
    )


def build_change_log_table():
    """构建变更记录表格（含 tblGrid）"""
    headers = ['更改时间', '更改人', '更改的主要内容', '批准人', '备注', '版本号']
    col_widths = [1800, 1200, 3000, 1200, 1200, 1200]

    tbl_xml = f'<w:tbl {nsdecls("w")}>'

    # 表格属性
    tbl_xml += '<w:tblPr>'
    tbl_xml += '<w:tblW w:w="9600" w:type="dxa"/>'
    tbl_xml += '<w:jc w:val="center"/>'
    tbl_xml += '<w:tblBorders>'
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tbl_xml += f'<w:{border} w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += '</w:tblBorders>'
    tbl_xml += '</w:tblPr>'

    # tblGrid（必须，否则 python-docx 读取会报错）
    tbl_xml += '<w:tblGrid>'
    for w in col_widths:
        tbl_xml += f'<w:gridCol w:w="{w}"/>'
    tbl_xml += '</w:tblGrid>'

    # 表头行（12pt 仿宋，不加粗，与参考文档一致）
    tbl_xml += '<w:tr>'
    for j, h in enumerate(headers):
        tbl_xml += (
            f'<w:tc>'
            f'<w:tcPr><w:tcW w:w="{col_widths[j]}" w:type="dxa"/></w:tcPr>'
            f'<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
            f'<w:r><w:rPr>'
            f'<w:rFonts w:eastAsia="仿宋" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
            f'</w:rPr><w:t>{h}</w:t></w:r></w:p></w:tc>'
        )
    tbl_xml += '</w:tr>'

    # 5 个空行
    for _ in range(5):
        tbl_xml += '<w:tr>'
        for j in range(len(headers)):
            tbl_xml += (
                f'<w:tc>'
                f'<w:tcPr><w:tcW w:w="{col_widths[j]}" w:type="dxa"/></w:tcPr>'
                f'<w:p/></w:tc>'
            )
        tbl_xml += '</w:tr>'

    tbl_xml += '</w:tbl>'
    return parse_xml(tbl_xml)


def remove_field_codes(doc: Document):
    """移除会触发"引用其他文件"弹窗的域代码。
    TOC 域在 sdt（结构化文档标签）内，需要特殊处理。"""
    body = doc.element.body
    external_keywords = {'INCLUDETEXT', 'LINK', 'INCLUDEPICTURE', 'IMPORT', 'DATABASE'}

    # 1. 处理 sdt 包裹的 TOC 域：解除 sdt 包装，保留 TOC 文本内容
    for sdt in list(body.findall(qn('w:sdt'))):
        sdt_content = sdt.find(qn('w:sdtContent'))
        if sdt_content is None:
            continue
        # 检查是否包含 TOC instrText
        has_toc = False
        for it in sdt_content.iter(qn('w:instrText')):
            if 'TOC' in (it.text or '').upper():
                has_toc = True
                break
        if not has_toc:
            continue

        # 先移除 sdt 内的域代码 run（fldChar 和 instrText 所在的 run）
        for r in list(sdt_content.iter(qn('w:r'))):
            if r.find(qn('w:fldChar')) is not None or r.find(qn('w:instrText')) is not None:
                r.getparent().remove(r)

        # 把 sdtContent 内的段落提升到 body 层级，替换 sdt
        parent = sdt.getparent()
        idx = list(parent).index(sdt)
        for child in list(sdt_content):
            parent.insert(idx, child)
            idx += 1
        parent.remove(sdt)

    # 处理 fldSimple（简单域）
    for fld in body.findall('.//' + qn('w:fldSimple')):
        instr = (fld.get(qn('w:instr')) or '').strip().upper()
        if any(kw in instr for kw in external_keywords):
            # 把域内容提升为普通 run，移除域包装
            parent = fld.getparent()
            idx = list(parent).index(fld)
            for child in list(fld):
                parent.insert(idx, child)
                idx += 1
            parent.remove(fld)

    # 处理 fldChar 复合域（BEGIN...SEPARATE...END）
    # 跨段落收集所有 body 内的 run 元素及其所属段落
    all_runs = []
    for para in body.iter(qn('w:p')):
        for r in para.findall(qn('w:r')):
            all_runs.append((r, para))

    i = 0
    while i < len(all_runs):
        r, para = all_runs[i]
        fld_char = r.find(qn('w:fldChar'))
        if fld_char is not None and fld_char.get(qn('w:fldCharType')) == 'begin':
            # 收集 instrText 直到 end
            instr_text = ''
            j = i + 1
            sep_idx = None
            while j < len(all_runs):
                rj, pj = all_runs[j]
                for it in rj.findall(qn('w:instrText')):
                    instr_text += (it.text or '')
                fc = rj.find(qn('w:fldChar'))
                if fc is not None:
                    if fc.get(qn('w:fldCharType')) == 'separate':
                        sep_idx = j
                    elif fc.get(qn('w:fldCharType')) == 'end':
                        break
                j += 1

            instr_upper = instr_text.upper().strip()
            should_remove = any(kw in instr_upper for kw in external_keywords)
            should_flatten = instr_upper.startswith('TOC ')

            if should_remove or should_flatten:
                # 移除 begin run
                all_runs[i][1].remove(all_runs[i][0])
                # 移除 instrText runs 和 separate run（begin+1 到 sep_idx）
                if sep_idx is not None:
                    for k in range(i + 1, sep_idx + 1):
                        try:
                            all_runs[k][1].remove(all_runs[k][0])
                        except ValueError:
                            pass
                # 移除 end run
                if j < len(all_runs):
                    try:
                        all_runs[j][1].remove(all_runs[j][0])
                    except ValueError:
                        pass
                # 如果是外部引用，连内容也删
                if should_remove and sep_idx is not None:
                    for k in range(sep_idx + 1, j):
                        try:
                            all_runs[k][1].remove(all_runs[k][0])
                        except ValueError:
                            pass
                # 重建 run 列表
                all_runs = []
                for para2 in body.iter(qn('w:p')):
                    for r2 in para2.findall(qn('w:r')):
                        all_runs.append((r2, para2))
                continue
        i += 1


def remove_pandoc_meta_blocks(doc: Document, meta: dict):
    """移除 pandoc 自动生成的 Title/Subtitle/Author/Date 段落
    pandoc 会把 YAML frontmatter 渲染到文档开头，但封面由后处理生成，所以需要清除"""
    meta_styles = {'Title', 'Subtitle', 'Author', 'Date'}
    meta_values = {v.strip() for v in meta.values() if v.strip()}
    body = doc.element.body
    # 只检查正文开头（第一个 Heading 之前）的段落
    for para in list(doc.paragraphs):
        if para.style and para.style.name.startswith('Heading'):
            break
        should_remove = (
            (para.style and para.style.name in meta_styles)
            or (para.text.strip() in meta_values)
        )
        if should_remove:
            body.remove(para._element)


def fix_chinese_fonts(doc: Document):
    """确保每个 run 的 eastAsia 字体正确"""
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        fonts = None
        for key in FONT_MAP:
            if style_name == key or style_name.startswith(key):
                fonts = FONT_MAP[key]
                break
        if not fonts:
            fonts = FONT_MAP['Normal']

        cn_font, en_font = fonts
        for run in para.runs:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rpr.insert(0, rfonts)
            rfonts.set(qn('w:eastAsia'), cn_font)
            if not rfonts.get(qn('w:ascii')):
                rfonts.set(qn('w:ascii'), en_font)
                rfonts.set(qn('w:hAnsi'), en_font)


def fix_table_formats(doc: Document):
    """修复所有表格：添加边框 + 按内容密度自动调整列宽"""
    border_xml = (
        '<w:tcBorders {ns}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    ).format(ns=nsdecls("w"))

    # 正文区域可用宽度（页宽 - 左右边距 = 21cm - 2.8cm - 2.6cm = 15.6cm ≈ 8840 twips）
    TABLE_WIDTH_TWIPS = 8840

    for table in doc.tables:
        num_cols = len(table.columns)
        if num_cols == 0:
            continue

        # ── 1. 按内容密度计算列宽 ──
        col_max_len = [0] * num_cols
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci < num_cols:
                    text_len = len(cell.text.strip())
                    col_max_len[ci] = max(col_max_len[ci], text_len)

        MIN_COL_TWIPS = 600   # 约 1cm

        # 表头长度
        header_lens = [len(cell.text.strip()) for cell in table.rows[0].cells]

        # 策略：
        # 1. 每列先获得基于表头的「底线宽度」（表头字符数 × 单字宽度）
        # 2. 剩余空间按内容密度分配，但单列不超过总宽 50%
        CHAR_WIDTH = 200  # 每个中文字符约 200 twips（14pt 字体）
        MAX_RATIO = 0.50  # 单列最大占比

        # 底线宽度：至少能放下表头（序号等短表头给最小宽度）
        base_widths = []
        for ci in range(num_cols):
            hl = header_lens[ci] if ci < len(header_lens) else 2
            base = max(hl * CHAR_WIDTH, MIN_COL_TWIPS)
            base_widths.append(base)

        base_total = sum(base_widths)

        if base_total >= TABLE_WIDTH_TWIPS:
            # 表头已经占满，按比例缩放
            col_widths = [int(TABLE_WIDTH_TWIPS * w / base_total) for w in base_widths]
        else:
            # 有剩余空间，按内容密度分配
            remaining = TABLE_WIDTH_TWIPS - base_total
            # 内容超出表头的部分作为额外需求
            extra_needs = []
            for ci in range(num_cols):
                content = col_max_len[ci]
                header = header_lens[ci] if ci < len(header_lens) else 2
                extra_needs.append(max(content - header, 0))

            extra_total = sum(extra_needs)
            col_widths = list(base_widths)

            if extra_total > 0:
                for ci in range(num_cols):
                    bonus = int(remaining * extra_needs[ci] / extra_total)
                    col_widths[ci] += bonus

            # 限制单列不超过 MAX_RATIO
            for ci in range(num_cols):
                cap = int(TABLE_WIDTH_TWIPS * MAX_RATIO)
                if col_widths[ci] > cap:
                    overflow = col_widths[ci] - cap
                    col_widths[ci] = cap
                    # 溢出部分均分给其他列
                    others = [j for j in range(num_cols) if j != ci]
                    per_other = overflow // len(others) if others else 0
                    for j in others:
                        col_widths[j] += per_other

        # 微调：确保总宽 = TABLE_WIDTH_TWIPS
        diff = TABLE_WIDTH_TWIPS - sum(col_widths)
        widest = col_widths.index(max(col_widths))
        col_widths[widest] += diff

        # ── 2. 设置表格宽度为固定值 ──
        tbl = table._tbl
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is not None:
            tw = tbl_pr.find(qn('w:tblW'))
            if tw is not None:
                tw.set(qn('w:w'), str(TABLE_WIDTH_TWIPS))
                tw.set(qn('w:type'), 'dxa')

        # ── 3. 更新 tblGrid ──
        grid = tbl.find(qn('w:tblGrid'))
        if grid is not None:
            grid_cols = grid.findall(qn('w:gridCol'))
            for ci, gc in enumerate(grid_cols):
                if ci < len(col_widths):
                    gc.set(qn('w:w'), str(col_widths[ci]))

        # ── 4. 更新每个单元格宽度 + 添加边框 ──
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                tc = cell._element
                tc_pr = tc.find(qn('w:tcPr'))
                if tc_pr is None:
                    tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")} />')
                    tc.insert(0, tc_pr)

                # 设置单元格宽度
                if ci < len(col_widths):
                    tc_w = tc_pr.find(qn('w:tcW'))
                    if tc_w is None:
                        tc_w = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[ci]}" w:type="dxa" />')
                        tc_pr.insert(0, tc_w)
                    else:
                        tc_w.set(qn('w:w'), str(col_widths[ci]))
                        tc_w.set(qn('w:type'), 'dxa')

                # 添加边框
                old_borders = tc_pr.find(qn('w:tcBorders'))
                if old_borders is not None:
                    tc_pr.remove(old_borders)
                tc_pr.append(parse_xml(border_xml))


def set_header(doc: Document, text: str):
    """设置页眉（所有节，右对齐，与参考文档一致）"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            p = header.paragraphs[0]
        else:
            p = header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')


def extract_meta(md_path: str) -> dict:
    """从 Markdown YAML frontmatter 提取元数据"""
    content = Path(md_path).read_text(encoding='utf-8')
    meta = {}
    match = re.match(r'^---\s*\n(.*?)\n---', content, re.DOTALL)
    if match:
        for line in match.group(1).split('\n'):
            if ':' in line:
                key, val = line.split(':', 1)
                meta[key.strip()] = val.strip().strip('"').strip("'")
    return meta


def insert_front_matter(doc: Document, meta: dict):
    """一次性按顺序插入：封面 → 分节符 → 变更记录 → 分节符"""
    title = meta.get('title', '项目名称')
    subtitle = meta.get('subtitle', '概要设计说明书')
    company = meta.get('company', '')
    date_str = meta.get('date', '')

    body = doc.element.body
    first_element = body[0] if len(body) > 0 else None

    # 按顺序构建所有前置元素
    elements = []

    # ── 封面（按参考文档结构：3空行→标题→副标题→12空行→公司→日期）──
    for _ in range(3):
        elements.append(make_para())

    # 标题（黑体 26pt）
    elements.append(make_para(title, cn_font='黑体', size=52))
    # 副标题
    elements.append(make_para(subtitle, cn_font='黑体', size=52))

    # 中间空行
    for _ in range(12):
        elements.append(make_para())

    # 公司名（仿宋 16pt）
    if company:
        elements.append(make_para(company, cn_font='仿宋', size=32))
    # 日期
    if date_str:
        elements.append(make_para(date_str, cn_font='仿宋', size=32))

    # 封面分节符
    elements.append(make_section_break())

    # ── 变更记录 ──
    elements.append(make_para())  # 空行
    elements.append(make_para('变更记录', cn_font='黑体', size=24, bold=True))
    elements.append(make_para())  # 空行
    elements.append(build_change_log_table())
    elements.append(make_para())  # 空行

    # 变更记录分节符（正文区域边距）
    elements.append(make_section_break(top=1985, bottom=1588, left=1588, right=1474))

    # 正序插入到文档开头（每次都在 first_element 之前插入）
    for elem in elements:
        if first_element is not None:
            idx = list(body).index(first_element)
            body.insert(idx, elem)
        else:
            body.append(elem)


def main():
    if len(sys.argv) < 3:
        print("用法: post_process.py <docx文件> <原始md文件>")
        sys.exit(1)

    docx_path = sys.argv[1]
    md_path = sys.argv[2]

    meta = extract_meta(md_path)
    doc = Document(docx_path)

    # 0. 清理域代码（消除打开弹窗）
    remove_field_codes(doc)

    # 0.1 移除 pandoc 自动生成的元数据段落（Title/Date 等）
    remove_pandoc_meta_blocks(doc, meta)

    # 1. 修正中文字体
    fix_chinese_fonts(doc)

    # 2. 插入封面 + 变更记录
    insert_front_matter(doc, meta)

    # 3. 表格边框（在插入变更记录表之后执行，确保所有表格都有边框）
    fix_table_formats(doc)

    # 4. 设置页眉
    header_text = meta.get('subtitle', meta.get('title', ''))
    if header_text:
        set_header(doc, header_text)

    doc.save(docx_path)
    print(f"后处理完成: {docx_path}")


if __name__ == '__main__':
    main()
